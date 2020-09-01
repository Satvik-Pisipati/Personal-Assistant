import pyttsx3 
import speech_recognition as sr
import datetime
import wikipedia 
import docx
import os
import time
import pyautogui as p
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from tkinter import*
window = Tk()
window.geometry("1000x800")
window.title("codexpy V_A")






def quit1():
    quit()

def window_check() : 

 engine = pyttsx3.init('sapi5')
 voices = engine.getProperty('voices')

 engine.setProperty('voice', voices[0].id)


 def speak(audio):
    engine.say(audio)
    engine.runAndWait()
 def wake():
    speak("can you order your god do something when he is not willing to do. no right because he is your creator,i am also in same situation like you Sir ,i am sorry ") 


 def gr():
    hour = int(datetime.datetime.now().hour)
    if hour>=0 and hour<12:
        speak("Good Morning!")

    elif hour>=12 and hour<18:
        speak("Good Afternoon!")   

    else:
        speak("Good Evening!")  

    speak("I am your assistant Sir. Please tell me how may I help you")
 def about():
    speak("I am June,a loyal personal computer butler. I am first pyaudio project.") 


 def takeCommand():
    

    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        r.pause_threshold = 1
        audio = r.listen(source)

    try:
        print("Recognizing...")    
        query = r.recognize_google(audio, language='en-in')
        print(f"User said: {query}\n")
        
        

    except Exception as e:
            
        speak("Say that again please...")  
        return "None"
    return query



 def profile() :
    sp= driver.find_element_by_xpath('//*[@id="side"]/header/div[1]/div').click()
    
    
    while True:
    
        query5 = takeCommand().lower()
        speak('what do u want')
        if 'view photo' in query5 :
            spv= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[1]/div/div/div/div/div/img').click()
            
            p.press('down')
            p.press('enter')
            speak('say close to close image')
            queryc = takeCommand().lower()
            if 'close' in queryc :
                sp= driver.find_element_by_xpath('//*[@id="app"]/div/span[3]/div/div/div[1]/div[2]/div/div/div').click()
                
            
        elif 'take photo' in query5 :
             spt= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[1]/div/div/div/div/div/img').click()
             p.press('down',presses=2)
             p.press('enter')
             queryca = takeCommand().lower()
             if 'capture' in queryca :
             
              spca= driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/span/div/div/div/div/div/div/div/div[2]/span/div/div/span').click()
              speak('its recommended to use your mouse , i cannot handle it')
        elif 'upload photo' in query5 :
             spup= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[1]/div/div/div/div/div/img').click()
             p.press('down',presses=3)
             p.press('enter')
             speak('sorry i cannot handle it please use your mouse for this')
        elif 'remove photo' in query5 :
             spr= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[1]/div/div/div/div/div/img').click()
             p.press('down',presses=4)
             p.press('enter')
             speak('are u sure??')
             queryre = takeCommand().lower()
             if 'yes' in queryre :
                    spr= driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/div/div/div/div/div/div[2]/div[2]').click()
             elif 'cancel' in queryre :
                    spr= driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/div/div/div/div/div/div[2]/div[1]').click()
                    
                    
             
        elif 'yes' in query5:
            cn= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[2]/div[2]/div[1]/span[2]/div/span').click()
            query6 = takeCommand().lower()
            p.press('ctrl','A')
            p.press('backspace')
            p.typewrite(query6)
            
            
            cn= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[2]/div[2]/div[1]/span[2]/div/span').click()
        elif 'about' in query5:
            ca= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[4]/div[2]/div[1]/span[2]/div/span').click()
            query7 = takeCommand().lower()
            press('ctrl','A')
            press('backspace')
            p.typewrite(query7)
            
            ca= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/div/div[4]/div[2]/div[1]/span[2]/div/span').click()
        elif 'back' in query5 :
            sb= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/div/header/div/div[1]/button/span ').click()
            break
        else :
            speak('sorry i didnot get u please try view,take photo,upload photo,remove photo,change name,about or say back to go to previous page')
            

  

 def multmsgpeople() :
    speak('what is the message??')
    
    queryml3=  takeCommand().lower()
    a=[]
    
    speak('whom you want to send, say add before telling their name, say done after saying their name')
    while True:
        querym1 =  takeCommand().lower()
        if 'yes' in querym1 :
            querym12 =  takeCommand().lower()
            a.append(queryml)
            time.sleep(2)
            del querym12

        
            
          
            
        elif 'done' in querym1 :
            break
    b= len(a)
    for i in range (0, len(a)+1):
        if i<=len(a):
            s4= driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').click()
            p.typewrite(a[i])
            p.hotkey('enter')
            s5 =  driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[2]')
            p.typewrite(queryml3)
            p.hotkey('enter')

        elif i>len(a) :
            speak('all messages are sent')
            break
    del queryml3
            
        



 def status() :

    ss= driver.find_element_by_xpath('//*[@id="side"]/header/div[2]/div/span/div[1]/div').click()
    speak(' sorry i cannot handle it use your mouse')
    speak('after viewing say assistant ok')
    while True :
        querykj = takeCommand().lower()
        if 'assistant ok' in querykj :
            speak('ok sir')
            break
        


 def menu() :
     ss= driver.find_element_by_xpath('//*[@id="side"]/header/div[2]/div/span/div[3]/div/span').click()
     while True:
    
        query8 = takeCommand().lower()
        if 'new group' in query8 :
            p.press('down')
            speak('say back , after going back say newgroup')
            
        elif 'create room' in query8 :
            
            speak('sorry i cannot handle it better use your mouse.say pause to pause me')
            p.press('down', presses =2)
            p.press('enter')
            
        elif 'profile' in query8 :
            p.press('down',presses = 3)
            p.press('enter')
            profile()
            
            break
        
        elif 'archived' in query8 :
            p.press('down',presses = 4)
            p.press('enter')
            speak('sorry i cannot handle it better use your mouse')
        elif 'starred' in query8 :
            speak('sorry i cannot handle it better use your mouse')
            p.press('enter')
            p.press('down',presses = 5)
        elif 'settings' in query8 :
            p.press('down',presses = 6)
            p.press('enter')
            while True:
    
             query9 = takeCommand().lower()
             speak('what do u want')
             if 'notifications' in query9 :
                     snnn= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/div[2]/div[2]/div/span').click()
                     
                     while True:
    
                       query10 = takeCommand().lower()
                       if 'sounds' in query10 :
                           snns= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/div/div[1]/div[2]').click()
                       elif 'desktop alert' in query10 :
                           snda= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/div/div[2]/div[2]').click()
                       elif 'show previews' in query10 :
                           snssp= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/div/div[3]/div[2]').click()
                       elif 'turn off all desktop notification' in query10 :
                           sntoa= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/div/div[4]/div/select').click()
                           speak('use ur mouse please')
                                 
                       elif 'back' in query10 :
                           snbppp = driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/header/div/div[1]/button/span').click()
                           break
                       else:
                           speak('try to say sounds,desktop alert,show previews or back to go to previous page')  
                           
                           
                     
             elif 'theme' in query9 :
                    snbthe = driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/div[3]/div[2]').click()
                    while True :
                     queryl = takeCommand().lower()
                     
                     #//*[@id="app"]/div/span[2]/div/div/div/div/div/div/div[2]/form/ol/li[1]/label/input
                     if 'ligth' in queryl :
                         snbthel = driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/div/div/div/div/div/div[2]/form/ol/li[1]/label/input').click()
                         p.press('enter')
                         break
                     elif 'dark' in queryl :
                         snbthed = driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/div/div/div/div/div/div[2]/form/ol/li[2]/label/input').click()
                         p.press('enter')
                         break
                     elif 'cancel' in queryl :
                         snbthecan = driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/div/div/div/div/div/div[3]/div[1]').click()
                         break
                     else :
                         speak('try to say light,dark or cancel to leave it')
                         
                         
             elif 'chat wallpaper' in query9 :
                     speak('use mouse for this operation ')
             elif 'blocked' in query9 :
                     speak('use mouse for this operation ')
             elif 'help' in query9 :
                     speak('use mouse for this operation ')
             elif 'back' in query9 :
                     snbppppp = driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/header/div/div[1]/button/span').click()
             else :
                     speak('sorry i didnot get u please try notifications,themes,chat wallpaper,blocked,help or say back to go to previous page')
        elif 'logout' in query8 :
            p.press('down',presses = 3)
            p.press('alt','f4')
            
        
        elif 'back' in query8 :
            ss= driver.find_element_by_xpath('//*[@id="side"]/header/div[2]/div/span/div[3]/div/span').click
            break
        else :
            speak('sorry i didnot get u please say new group , create room ,profile ,archieved ,starred settings, logout or say back to go to previous page')

            



 def new_group() :
     ss= driver.find_element_by_xpath('//*[@id="side"]/header/div[2]/div/span/div[3]/div/span').click()
     
     while True:
    
        query8 = takeCommand().lower()
        if 'new group' in query8 :
           
            speak('whom u want to add')
            speak('say yes before adding them')
            while True:
             queryadd1 = takeCommand().lower()
             if 'add' in queryadd1 :
              
              while True :
               
            
               queryadd = takeCommand().lower()
               p.typewrite(queryadd )
               p.press('down')
               del queryadd
             
               while True:
                 speak('this one')
                 queryaddu = takeCommand().lower()
                 if 'yes' in queryaddu :
                     p.press('enter')
                     break
                 elif 'no' in queryaddu :
                     p.press('down')
                 elif 'next' in queryaddu :
                     p.press('down')
                 elif 'up' in queryaddu :
                     p.press('up')
                 elif 'wrong' in queryaddu :
                     speak('ok')
                     sclos= driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/span/button').click()
                     
               speak('make group')
               del queryaddu 
             elif 'make group' in queryadd1 :
              mg= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/span/div/span').click()
              while True :
               speak('say yes')
               
               queryadduna = takeCommand().lower() 
               if 'yes' in queryadduna :
                 p.press('ctrl','A')
                 p.press('backspace')
                 while True :
                     queryaddunai = takeCommand().lower()
                     speak('say note and tell ur group name')
                     if 'note' in queryaddunai :
                         while True :
                             queryaddi = takeCommand().lower()
                             p.typewrite(queryaddi)
                             del queryaddi
                             speak('if u want to change name then say yes or say no')
                             
                             break
                     elif 'yes' in  queryaddunai :
                         speak('say yes')
                         break
                     elif 'no' in queryaddunai :
                         mgf= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/div/span/div/div/span').click()
                         speak('say ok')
                         break
               elif 'ok' in  queryadduna :
                   break
               elif 'back' in queryadduna:
                   mgf= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/header/div/div[1]/button/span').click()
                   break
            break         
        elif 'back' in  query8 :
            mgff= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[1]/span/div/span/div/header/div/div[1]/button/span').click()


 def media() :
    mgfsssf= driver.find_element_by_xpath('//*[@id="main"]/header/div[3]/div/div[3]/div').click()
    p.hotkey('down')
    p.hotkey('enter')
    mgfsssyf= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[3]/span/div/span/div/div/div[1]/div[2]/div[1]/div').click()
    time.sleep(3)
    
    mgfsssyf= driver.find_element_by_xpath('//*[@id="app"]/div/div/div[2]/div[3]/span/div/span/div/div[2]/span/div/div/div/div[3]').click()
    
    
    
    speak('say next to view next photo or previous to view previous photo')
    while True :
      query2aa = takeCommand().lower()
      if 'next' in query2aa :
        p.press('left')
      elif 'previous' in query2aa :
        p.press('right')
      elif 'star' in query2aa :
       mgfsssysf= driver.find_element_by_xpath('//*[@id="app"]/div/span[3]/div/div/div[2]/div[1]/div[2]/div/div[2]/div/span').click()
      elif 'download' in query2aa :
        mgfsssysf= driver.find_element_by_xpath('//*[@id="app"]/div/span[3]/div/div/div[2]/div[1]/div[2]/div/div[4]/div/span').click()
      elif 'forward' in query2aa :
         mgfsssysf= driver.find_element_by_xpath('//*[@id="app"]/div/span[3]/div/div/div[2]/div[1]/div[2]/div/div[3]/div/span').click()
         while True:
                   speak('say yes to add names')
                   querycf = takeCommand().lower()
                   if 'yes' in query2 :
                      speak('to whom ')
                      
                      
                      queryaaaaa = takeCommand().lower()
                      saa= driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').click()
                      
                      p.typewrite(query4)
                      time.sleep(5)
                      speak('say yes to accept,no or next for next, up ,wrong change search')
                      p.press('down')
                      speak('this one ??')
                      
                      while True :
                          querybjaa = takeCommand().lower()
                          if 'no' in querybjaa:
                           p.press('down')
                          elif 'next' in querybjaa :
                           p.press('down')
                          elif 'up' in querybjaa :
                           p.press('up')
                          elif 'wrong' in querybjaa :
                            speak('ok')
                            sclos= driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/span/button').click()
                            
                            del querybjaa
                            break
                          elif 'yes' in querybjaa :
                            p.press('enter')
                    
                   elif 'ok' in querybjaa :
                         saa= driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/span/div/div/div/div/div/div/span/div/div/div').click()
                         
                         break
                   elif 'close' in querybjaa :
                       saa= driver.find_element_by_xpath('//*[@id="app"]/div/span[2]/div/span/div/div/div/div/div/div/header/div/div[1]/button').click()
                       break
                   else :
                    speak('try to say  no,next,up,wrong,close')
                       
                       
      elif 'close' in query2aa :
         mgfsssysf= driver.find_element_by_xpath('//*[@id="app"]/div/span[3]/div/div/div[2]/div[1]/div[2]/div/div[5]/div').click()
      elif 'quit' in query2aa :
          quit()
      else :
        speak('try to say  next,previous,star,download,forward,close')
   
       

    
 def msg_box() :
                  
                
                  while True:
                   speak('say yes to message someone ')
                   query2 = takeCommand().lower()
                   if 'yes' in query2 :
                      
                      
                      
                      
                      s4= driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]').click()
                      
                      
                      speak('to whom ')
                      query4 = takeCommand().lower()
                      p.typewrite(query4)
                      time.sleep(5)
                      speak('say yes to accept,no or next for next, up ,wrong change search')
                      
                      p.press('down')
                      
                      while True :
                         
                          speak('this one ??')
                          querybj = takeCommand().lower()
                          if 'no' in querybj:
                           p.press('down')
                          elif 'next' in querybj :
                           p.press('down')
                          elif 'up' in querybj :
                           p.press('up')
                          elif 'wrong' in querybj :
                            speak('ok')
                            sclos= driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/span/button').click()
                            
                            del querybj
                            break
                          elif 'yes' in querybj :
                            p.press('enter')
                            
                            while True :
                             speak('say yes to start messaging,media for media,ok stop messaging')
                             
                              
                             query3 = takeCommand().lower()
                             if 'yes' in query3:
                               del query3 
                               speak('whats the message')
                               queryz = takeCommand().lower()
                               
                               s5 =  driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[2]')
                               s5.send_keys(queryz)
                               p.press('enter')
                               speak('if u want to stop messaging say ok or yes to continue')
                               del queryz
                               
                             elif 'media' in query3 :
                                 media()
                             elif 'ok' in query3:
                              break
                            break
                          elif 'quit' in  query :
                                 quit()
                          else:
                              speak('say yes to accept,no or next for next, up ,wrong change search')
                           
                                 
                                          
                                        
                      del query4 
                      del query2
                   
                   elif 'ok' in query2 :
                       break
                   elif 'quit' in  query :
                      quit()
                   else :
                       speak('say yes to open message someone or ok to stop ')
                   

            
            
            
            
    


 if __name__ == "__main__":
    gr()
    while True:
    
        query = takeCommand().lower()
        speak('say open')
        if 'open' in query:
            
            driver = webdriver.Chrome(r"C:\Users\asgav\Desktop\crtrdrtdtrdt\chromedriver.exe")
            driver.get("https://web.whatsapp.com/")
            time.sleep(5)
            speak('ready')

            while True:
    
              query1 = takeCommand().lower()
              if 'profile' in query1 :
                  profile()
              elif 'status' in query1 :
                  status()
              elif 'new group' in query1 :
                  new_group()
              elif 'menu' in query1 :
                  menu()
              elif 'message' in query1 :###search engine and message chat area--
                  msg_box()
              elif 'yes' in query1 :
                   multmsgpeople()
              elif 'quit' in  query :
                  quit()
               
              else :


                  speak('say profile, status, new group, menu, message')

        elif 'wikipedia' in query:
                 speak('Searching Wikipedia...')
                 query = query.replace("wikipedia", "")
                 results = wikipedia.summary(query, sentences=2)
                 speak("According to Wikipedia")
                 speak(results)
        elif 'take notes' in query :
            speak('how should i save this')
            queryqq = takeCommand().lower()
            a =queryqq+'.docx'
            doc = docx.Document(a)
            while True :
                queryqq1 = takeCommand().lower()
                speak('say yes to start writing paragraph')
                if 'yes' in queryqq1 :
                   queryqq2 = takeCommand().lower()
                   doc = docx.Document()
                   doc.add_paragraph(queryqq12).text
                elif 'code save' in queryqq1 :
                    doc.save(a)
                    break
                    
            
            
        elif 'quit' in  query :
            quit()
            
       
label1 = Label(window,text = "VOICE ASSISTANT  BY CODE_BRO'S\n\n\nfor whatsappweb say open\n\n\nfor wikipedia say wikipedia\n\n\nfor time say time\n\n\nwikipedia\n\n\ntake notes",fg='black',bg ='white',relief="solid",font=("arial",21,"bold"))
label1.pack(fill=BOTH,pady=2,padx=2,expand=True)


button1=Button(window,text="start V_A",fg='blue',bg='white',relief=RIDGE,font=("arial",12,"bold"),command=window_check)
button1.place(x=300,y=700)
button2=Button(window,text="quit",fg='red',bg='white',relief=RIDGE,font=("arial",12,"bold"),command=quit1)
button2.place(x=700,y=700)

        
            



        
